"""The ``ats_plain`` template plugin — a Word document built for machines to read.

This is the maximum-compatibility option, and the one to reach for when an application form
says "upload your resume (.doc, .docx)" or when a particular ATS is known to mangle PDFs.

Why a DOCX at all, when the LaTeX path produces something far better looking? Because PDF
text extraction is a heuristic. A PDF stores glyphs at coordinates; recovering "these glyphs
were a bulleted line under the heading *Experience*" is inference, and older parsers get it
wrong in ways that silently drop half a resume. A ``.docx`` needs no inference: the bullet is
a ``List Bullet`` paragraph, the heading is a ``Heading 1`` paragraph, and every parser in
existence reads the same structure the author wrote.

The document is built programmatically rather than from a template file, because the
structure *is* the product here — there is no visual layer to separate out.

Everything that makes a resume parser fail is deliberately absent:

* **no tables.** The single most common cause of scrambled parses. The right-aligned date
  column is a right tab stop inside an ordinary paragraph, which is one text run in document
  order, not two cells.
* **no text boxes, no shapes, no images, no headers or footers.** Content in a text box lives
  outside the main document body and is routinely skipped entirely.
* **no multi-column section layout.**
* **standard fonts only** — Calibri by default, Arial on request. Both are metric-safe
  everywhere, so the document paginates the same on the reviewer's machine as on ours.
* **plain-text URLs.** ``github.com/jane`` is printed as text rather than as a hyperlink
  field: the visible string is then identical to the extracted string, and there is no
  relationship indirection for a parser to resolve. A human can still copy it.

PDF conversion goes through LibreOffice, which is the only faithful DOCX renderer that can be
driven headlessly. When it is absent the render fails loudly — returning a ``.docx`` from a
call that asked for a ``.pdf`` would send the wrong file to an employer.
"""

from __future__ import annotations

import asyncio
import shutil
import tempfile
from pathlib import Path
from typing import Any, ClassVar, Final

import structlog

from app.documents.markdown import (
    CONTACT_SEPARATOR,
    DEFAULT_CLOSING,
    DEFAULT_SKILLS_HEADING,
    DEFAULT_SUMMARY_HEADING,
    SKILLS_HEADING_META_KEY,
    SUMMARY_HEADING_META_KEY,
    RenderSettings,
    contact_items,
    estimated_page_count,
    generated_heading,
    recipient_lines,
    require_format,
    resolve_render_settings,
)
from app.documents.models import CoverLetterDocument, ResumeDocument, ResumeEntry
from app.documents.renderer import DocumentRenderError, RenderResult, TemplatePlugin
from app.models.enums import PluginKind
from app.plugins import PluginMeta, plugin

__all__ = [
    "DOCX_FORMAT",
    "SOFFICE_CANDIDATES",
    "SOFFICE_TIMEOUT_SECONDS",
    "AtsPlainTemplate",
    "docx_to_pdf",
    "resolve_soffice",
]

logger = structlog.get_logger(__name__)


# ======================================================================================
# Constants
# ======================================================================================

#: Formats this plugin produces.
DOCX_FORMAT: Final[str] = "docx"
PDF_FORMAT: Final[str] = "pdf"
DOCX_FORMATS: Final[frozenset[str]] = frozenset({DOCX_FORMAT, PDF_FORMAT})

#: Reported as ``RenderResult.engine``.
DOCX_ENGINE: Final[str] = "python-docx"
SOFFICE_ENGINE: Final[str] = "libreoffice"

#: Page margin used when the caller specifies none. Half an inch is the tightest margin that
#: every consumer printer reproduces without clipping, and it buys roughly four extra lines
#: over Word's one-inch default.
DEFAULT_DOCX_MARGIN_IN: Final[float] = 0.5

#: Fonts offered through ``options["font"]``. Restricted to two metric-stable faces that ship
#: with every Windows and Office install; a font the reviewer does not have would be
#: substituted and would repaginate the document.
SUPPORTED_FONTS: Final[tuple[str, ...]] = ("Calibri", "Arial")
DEFAULT_FONT: Final[str] = "Calibri"

#: Type scale, as multiples of the body size. Mirrors the LaTeX templates so the same resume
#: rendered both ways reads as the same document.
NAME_SIZE_RATIO: Final[float] = 1.95
HEADING_SIZE_RATIO: Final[float] = 1.08
CONTACT_SIZE_RATIO: Final[float] = 0.95

#: Vertical rhythm, in points.
SPACE_BEFORE_HEADING_PT: Final[float] = 9.0
SPACE_AFTER_HEADING_PT: Final[float] = 2.0
SPACE_BEFORE_ENTRY_PT: Final[float] = 5.0
BODY_LINE_SPACING: Final[float] = 1.06

#: Indent applied to bullet paragraphs, in inches.
BULLET_INDENT_IN: Final[float] = 0.22

#: Colour of the rule under a section heading, as an OOXML hex triplet.
HEADING_RULE_COLOR: Final[str] = "9AA3AD"

#: Rule thickness in eighths of a point, which is the unit ``w:sz`` uses.
HEADING_RULE_SIZE_EIGHTHS: Final[str] = "6"

#: Elements that must follow ``w:pBdr`` inside ``w:pPr``. OOXML's ``CT_PPr`` is a sequence,
#: not a bag, so a border appended at the end validates in Word but is rejected by stricter
#: readers — including some ATS ingestion pipelines. python-docx models none of these
#: children, so the insertion point is spelled out here.
_PPR_AFTER_PBDR: Final[tuple[str, ...]] = (
    "w:shd",
    "w:tabs",
    "w:suppressAutoHyphens",
    "w:kinsoku",
    "w:wordWrap",
    "w:overflowPunct",
    "w:topLinePunct",
    "w:autoSpaceDE",
    "w:autoSpaceDN",
    "w:bidi",
    "w:adjustRightInd",
    "w:snapToGrid",
    "w:spacing",
    "w:ind",
    "w:contextualSpacing",
    "w:jc",
    "w:outlineLvl",
    "w:rPr",
    "w:sectPr",
)

#: Executables tried, in order, when converting to PDF.
SOFFICE_CANDIDATES: Final[tuple[str, ...]] = ("soffice", "libreoffice")

#: Absolute paths checked when neither name is on ``PATH``. LibreOffice's Windows and macOS
#: installers do not add themselves to ``PATH``, so "installed but not found" would otherwise
#: be the common case rather than the rare one.
SOFFICE_FALLBACK_PATHS: Final[tuple[str, ...]] = (
    r"C:\Program Files\LibreOffice\program\soffice.exe",
    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    "/usr/bin/soffice",
    "/usr/bin/libreoffice",
    "/usr/lib/libreoffice/program/soffice",
    "/snap/bin/libreoffice",
)

#: Wall-clock budget for one conversion. LibreOffice's first headless start builds a user
#: profile, which is the slow case this number is sized for.
SOFFICE_TIMEOUT_SECONDS: Final[float] = 180.0

#: Style names taken from the default python-docx template.
STYLE_TITLE: Final[str] = "Title"
STYLE_HEADING: Final[str] = "Heading 1"
STYLE_BULLET: Final[str] = "List Bullet"
STYLE_NORMAL: Final[str] = "Normal"


# ======================================================================================
# LibreOffice conversion
# ======================================================================================


def resolve_soffice() -> str:
    """Find the LibreOffice executable.

    Returns:
        An absolute path to ``soffice``.

    Raises:
        DocumentRenderError: If LibreOffice is not installed. The message names the install
            routes and the two alternatives that need no system binary — asking for ``docx``
            instead of ``pdf``, or switching to the ``web`` template, whose reportlab fallback
            is pure Python.
    """
    for candidate in SOFFICE_CANDIDATES:
        found = shutil.which(candidate)
        if found:
            return found
    for path in SOFFICE_FALLBACK_PATHS:
        candidate_path = Path(path)
        if candidate_path.is_file():
            return str(candidate_path)

    raise DocumentRenderError(
        "LibreOffice is required to convert a .docx to PDF and was not found. Install it "
        "(https://www.libreoffice.org/download — `winget install LibreOffice.LibreOffice`, "
        "`brew install --cask libreoffice`, or `apt install libreoffice-writer`), or put "
        "soffice on PATH. To get a PDF without it, render with the 'web' template, whose "
        "reportlab fallback is pure Python, or with 'modern'/'classic' if a LaTeX engine is "
        "available. Requesting fmt='docx' from this template needs no external binary at all."
    )


async def docx_to_pdf(
    path: Path,
    out: Path | None = None,
    *,
    timeout: float = SOFFICE_TIMEOUT_SECONDS,
) -> Path:
    """Convert a ``.docx`` to PDF with headless LibreOffice.

    Args:
        path: The source ``.docx``. Must exist.
        out: Destination PDF. Defaults to *path* with a ``.pdf`` suffix.
        timeout: Wall-clock budget in seconds.

    Returns:
        The path to the PDF.

    Raises:
        DocumentRenderError: If LibreOffice is absent, cannot be started, times out, exits
            non-zero, or exits zero without producing a PDF. The failure is never swallowed:
            a caller that asked for a PDF and quietly received a ``.docx`` would upload the
            wrong file type to an employer's portal.
    """
    if not path.is_file():
        raise DocumentRenderError(f"cannot convert {path}: file does not exist")

    executable = resolve_soffice()
    destination = out if out is not None else path.with_suffix(".pdf")

    with tempfile.TemporaryDirectory(prefix="applicantos-soffice-") as tmp:
        workdir = Path(tmp)
        # A private user profile: a headless run shares the profile with any interactive
        # LibreOffice the user has open, and the second process exits immediately without
        # converting anything. This is the single most common cause of "it works on my
        # machine" in this code path.
        profile = workdir / "profile"
        profile.mkdir()
        outdir = workdir / "out"
        outdir.mkdir()

        argv = [
            executable,
            f"-env:UserInstallation={profile.as_uri()}",
            "--headless",
            "--norestore",
            "--invisible",
            "--convert-to",
            "pdf",
            "--outdir",
            str(outdir),
            str(path),
        ]

        try:
            process = await asyncio.create_subprocess_exec(
                *argv,
                cwd=str(workdir),
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
            )
        except OSError as exc:
            raise DocumentRenderError(f"could not start LibreOffice ({executable}): {exc}") from exc

        try:
            stdout, stderr = await asyncio.wait_for(process.communicate(), timeout=timeout)
        except TimeoutError as exc:
            process.kill()
            await process.wait()
            raise DocumentRenderError(
                f"LibreOffice timed out after {timeout:.0f}s converting {path.name} to PDF."
            ) from exc

        output = (
            stdout.decode("utf-8", errors="replace") + stderr.decode("utf-8", errors="replace")
        ).strip()
        if process.returncode:
            raise DocumentRenderError(
                f"LibreOffice exited with code {process.returncode} converting "
                f"{path.name} to PDF.\n{output}"
            )

        produced = outdir / f"{path.stem}.pdf"
        if not produced.is_file():
            raise DocumentRenderError(
                f"LibreOffice reported success but produced no PDF for {path.name}.\n{output}"
            )

        try:
            destination.parent.mkdir(parents=True, exist_ok=True)
            shutil.copyfile(produced, destination)
        except OSError as exc:
            raise DocumentRenderError(f"could not write {destination}: {exc}") from exc

    logger.info("docx.converted", source=str(path), path=str(destination))
    return destination


# ======================================================================================
# Document construction
# ======================================================================================


def _resolve_font(options: dict[str, Any] | None) -> str:
    """Return the typeface to build the document in.

    Args:
        options: The caller's ``options`` mapping.

    Returns:
        A name from :data:`SUPPORTED_FONTS`.

    Raises:
        DocumentRenderError: If ``options["font"]`` names an unsupported face. Substituting
            silently would produce a document that repaginates on the reviewer's machine,
            which defeats the point of this template.
    """
    requested = (options or {}).get("font")
    if requested is None:
        return DEFAULT_FONT
    for supported in SUPPORTED_FONTS:
        if isinstance(requested, str) and requested.strip().lower() == supported.lower():
            return supported
    raise DocumentRenderError(
        f"font {requested!r} is not supported by the ats_plain template; "
        f"choose one of {list(SUPPORTED_FONTS)}"
    )


def _apply_font(run: Any, *, font: str, size_pt: float, bold: bool = False) -> None:
    """Set a run's typeface explicitly, including the East-Asian and complex-script slots.

    Word resolves a run's font from three separate attributes (``w:ascii``, ``w:eastAsia``,
    ``w:cs``); setting only the first leaves the others inherited, and a document that then
    passes through a non-English Word build can come back in a substituted face.

    Args:
        run: A ``docx.text.run.Run``.
        font: Typeface name.
        size_pt: Size in points.
        bold: Whether the run is bold.
    """
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), font)
    rfonts.set(qn("w:cs"), font)


def _add_bottom_border(paragraph: Any) -> None:
    """Draw a hairline rule under *paragraph*, matching the LaTeX section rule.

    A paragraph border, not a table and not a drawn shape: it is a property of the paragraph
    the heading already occupies, so it adds nothing to the text stream a parser reads.

    Args:
        paragraph: A ``docx.text.paragraph.Paragraph``.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), HEADING_RULE_SIZE_EIGHTHS)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), HEADING_RULE_COLOR)
    borders.append(bottom)

    properties = paragraph._p.get_or_add_pPr()
    properties.insert_element_before(borders, *_PPR_AFTER_PBDR)


def _spacing(
    paragraph: Any,
    *,
    before_pt: float = 0.0,
    after_pt: float = 0.0,
    line: float | None = None,
) -> None:
    """Set a paragraph's vertical spacing.

    Args:
        paragraph: A ``docx.text.paragraph.Paragraph``.
        before_pt: Space above, in points.
        after_pt: Space below, in points.
        line: Line spacing multiple, or ``None`` to leave it inherited.
    """
    from docx.shared import Pt

    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before_pt)
    fmt.space_after = Pt(after_pt)
    if line is not None:
        fmt.line_spacing = line


def _configure_document(document: Any, *, font: str, settings: RenderSettings) -> float:
    """Apply page geometry and the ``Normal`` style to a fresh document.

    Args:
        document: A ``docx.document.Document``.
        font: Typeface for the whole document.
        settings: Resolved font size and margin.

    Returns:
        The printable width in inches, which is where the right tab stop goes.
    """
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    normal = document.styles[STYLE_NORMAL]
    normal.font.name = font
    normal.font.size = Pt(settings.font_size)
    normal_rpr = normal.element.get_or_add_rPr()
    normal_rfonts = normal_rpr.get_or_add_rFonts()
    normal_rfonts.set(qn("w:eastAsia"), font)
    normal_rfonts.set(qn("w:cs"), font)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = BODY_LINE_SPACING

    for section in document.sections:
        section.top_margin = Inches(settings.margin_in)
        section.bottom_margin = Inches(settings.margin_in)
        section.left_margin = Inches(settings.margin_in)
        section.right_margin = Inches(settings.margin_in)

    return settings.content_width_in


def _add_heading(document: Any, text: str, *, font: str, settings: RenderSettings) -> None:
    """Add a section heading with a rule under it.

    The ``Heading 1`` *style* is kept — that outline level is what makes a parser treat the
    following paragraphs as a section — while the visual properties are overridden, because
    the stock style is blue Calibri Light and this template is black text only.

    Args:
        document: The document under construction.
        text: Heading text.
        font: Typeface.
        settings: Resolved font size and margin.
    """
    paragraph = document.add_heading("", level=1)
    run = paragraph.add_run(text)
    _apply_font(run, font=font, size_pt=settings.font_size * HEADING_SIZE_RATIO, bold=True)
    _spacing(
        paragraph,
        before_pt=SPACE_BEFORE_HEADING_PT,
        after_pt=SPACE_AFTER_HEADING_PT,
        line=1.0,
    )
    paragraph.paragraph_format.keep_with_next = True
    _add_bottom_border(paragraph)


def _add_body(
    document: Any,
    text: str,
    *,
    font: str,
    settings: RenderSettings,
    before_pt: float = 0.0,
) -> None:
    """Add one paragraph of body text.

    Args:
        document: The document under construction.
        text: The paragraph text.
        font: Typeface.
        settings: Resolved font size and margin.
        before_pt: Space above, in points.
    """
    paragraph = document.add_paragraph()
    _apply_font(paragraph.add_run(text), font=font, size_pt=settings.font_size)
    _spacing(paragraph, before_pt=before_pt, line=BODY_LINE_SPACING)


def _add_entry(
    document: Any,
    entry: ResumeEntry,
    *,
    font: str,
    settings: RenderSettings,
    tab_width_in: float,
    first: bool,
) -> None:
    """Add one resume entry: a header line with a right-aligned date, then its bullets.

    The header is a single paragraph containing ``Title, Organization`` + TAB + ``Dates``,
    with one right-aligned tab stop at the right margin. That is the DOCX equivalent of
    LaTeX's ``\\hfill`` — one paragraph, one text stream, correct reading order — where a
    two-cell table would put the date in a separate cell that some parsers read as the start
    of a new record.

    Args:
        document: The document under construction.
        entry: The entry to add.
        font: Typeface.
        settings: Resolved font size and margin.
        tab_width_in: Printable width, where the right tab stop is placed.
        first: Whether this is the first entry of its section (tighter space above).
    """
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.shared import Inches

    header = document.add_paragraph()
    _spacing(
        header,
        before_pt=0.0 if first else SPACE_BEFORE_ENTRY_PT,
        line=BODY_LINE_SPACING,
    )
    header.paragraph_format.tab_stops.add_tab_stop(Inches(tab_width_in), WD_TAB_ALIGNMENT.RIGHT)
    header.paragraph_format.keep_with_next = True

    if entry.title:
        _apply_font(header.add_run(entry.title), font=font, size_pt=settings.font_size, bold=True)
    if entry.title and entry.organization:
        _apply_font(header.add_run(", "), font=font, size_pt=settings.font_size)
    if entry.organization:
        _apply_font(header.add_run(entry.organization), font=font, size_pt=settings.font_size)

    right = CONTACT_SEPARATOR.join(part for part in (entry.location, entry.date_range) if part)
    if right:
        _apply_font(header.add_run(f"\t{right}"), font=font, size_pt=settings.font_size)

    for bullet in entry.bullets:
        if not bullet.strip():
            continue
        paragraph = document.add_paragraph(style=STYLE_BULLET)
        _apply_font(paragraph.add_run(bullet), font=font, size_pt=settings.font_size)
        _spacing(paragraph, line=BODY_LINE_SPACING)
        paragraph.paragraph_format.left_indent = Inches(BULLET_INDENT_IN)
        paragraph.paragraph_format.first_line_indent = Inches(-BULLET_INDENT_IN * 0.5)


def _import_docx() -> Any:
    """Import python-docx, with an actionable error when it is missing.

    Returns:
        The ``docx`` module.

    Raises:
        DocumentRenderError: If python-docx is not installed.
    """
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - exercised only without python-docx
        raise DocumentRenderError(
            "python-docx is required by the ats_plain template. Install it with "
            "`pip install python-docx`, or render with the 'web' or 'markdown' template."
        ) from exc
    return docx


def build_resume_docx(
    doc: ResumeDocument,
    *,
    font: str,
    settings: RenderSettings,
) -> Any:
    """Build the Word document for a resume.

    No escaping happens anywhere in this module, and none is needed: python-docx writes text
    into XML text nodes through lxml, which escapes ``&``, ``<`` and ``>`` itself. Passing
    LaTeX-escaped strings in here would print literal backslashes.

    Args:
        doc: The resume to render, exactly as the resume engine produced it.
        font: Typeface for the whole document.
        settings: Resolved font size and margin.

    Returns:
        A ``docx.document.Document`` ready to save.

    Raises:
        DocumentRenderError: If python-docx is not installed.
    """
    docx = _import_docx()
    document = docx.Document()
    tab_width_in = _configure_document(document, font=font, settings=settings)

    if doc.contact.name:
        title = document.add_paragraph(style=STYLE_TITLE)
        _apply_font(
            title.add_run(doc.contact.name),
            font=font,
            size_pt=settings.font_size * NAME_SIZE_RATIO,
            bold=True,
        )
        _spacing(title, after_pt=1.0, line=1.0)

    items = contact_items(doc.contact)
    if items:
        contact = document.add_paragraph()
        _apply_font(
            contact.add_run(CONTACT_SEPARATOR.join(item.text for item in items)),
            font=font,
            size_pt=settings.font_size * CONTACT_SIZE_RATIO,
        )
        _spacing(contact, line=BODY_LINE_SPACING)

    if doc.summary.strip():
        _add_heading(
            document,
            generated_heading(doc, SUMMARY_HEADING_META_KEY, DEFAULT_SUMMARY_HEADING),
            font=font,
            settings=settings,
        )
        _add_body(document, doc.summary, font=font, settings=settings)

    for section in doc.sections:
        if section.heading:
            _add_heading(document, section.heading, font=font, settings=settings)
        for index, entry in enumerate(section.entries):
            _add_entry(
                document,
                entry,
                font=font,
                settings=settings,
                tab_width_in=tab_width_in,
                first=index == 0,
            )

    if doc.skills_line.strip():
        _add_heading(
            document,
            generated_heading(doc, SKILLS_HEADING_META_KEY, DEFAULT_SKILLS_HEADING),
            font=font,
            settings=settings,
        )
        _add_body(document, doc.skills_line, font=font, settings=settings)

    return document


def build_cover_letter_docx(
    letter: CoverLetterDocument,
    *,
    font: str,
    settings: RenderSettings,
    closing: str,
) -> Any:
    """Build the Word document for a cover letter.

    Args:
        letter: The letter to render.
        font: Typeface for the whole document.
        settings: Resolved font size and margin.
        closing: Sign-off line.

    Returns:
        A ``docx.document.Document`` ready to save.

    Raises:
        DocumentRenderError: If python-docx is not installed.
    """
    docx = _import_docx()
    document = docx.Document()
    _configure_document(document, font=font, settings=settings)

    contact = letter.contact
    if contact.name:
        title = document.add_paragraph(style=STYLE_TITLE)
        _apply_font(
            title.add_run(contact.name),
            font=font,
            size_pt=settings.font_size * NAME_SIZE_RATIO * 0.8,
            bold=True,
        )
        _spacing(title, after_pt=1.0, line=1.0)

    items = contact_items(contact)
    if items:
        line = document.add_paragraph()
        _apply_font(
            line.add_run(CONTACT_SEPARATOR.join(item.text for item in items)),
            font=font,
            size_pt=settings.font_size * CONTACT_SIZE_RATIO,
        )
        _spacing(line, after_pt=SPACE_BEFORE_HEADING_PT, line=BODY_LINE_SPACING)

    if letter.date.strip():
        _add_body(document, letter.date, font=font, settings=settings, before_pt=6.0)

    for index, address_line in enumerate(recipient_lines(letter)):
        _add_body(
            document,
            address_line,
            font=font,
            settings=settings,
            before_pt=SPACE_BEFORE_HEADING_PT if index == 0 else 0.0,
        )

    _add_body(document, letter.salutation(), font=font, settings=settings, before_pt=12.0)
    for paragraph in letter.paragraphs():
        _add_body(document, paragraph, font=font, settings=settings, before_pt=8.0)

    _add_body(document, closing, font=font, settings=settings, before_pt=12.0)
    if contact.name:
        _add_body(document, contact.name, font=font, settings=settings, before_pt=12.0)

    return document


def _save(document: Any, out: Path) -> int:
    """Save a Word document, creating parent directories.

    Args:
        document: The document to save.
        out: Destination path.

    Returns:
        The size of the written file in bytes.

    Raises:
        DocumentRenderError: If the file cannot be written.
    """
    try:
        out.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(out))
    except OSError as exc:
        raise DocumentRenderError(f"could not write {out}: {exc}") from exc
    return out.stat().st_size


# ======================================================================================
# Plugin
# ======================================================================================


@plugin
class AtsPlainTemplate(TemplatePlugin):
    """The ``ats_plain`` template (``PluginKind.TEMPLATE``) — Word, built for parsers.

    Produces ``.docx`` natively and ``.pdf`` through LibreOffice. Nothing about the output is
    decorative: no tables, no text boxes, no images, no headers or footers, standard fonts,
    half-inch margins, and structural paragraph styles a resume parser can read without
    guessing.
    """

    meta: ClassVar[PluginMeta] = PluginMeta(
        kind=PluginKind.TEMPLATE,
        name="ats_plain",
        version="1.0.0",
        display_name="ATS Plain (Word)",
        description="Maximum-compatibility .docx: no tables, no boxes, standard fonts.",
        capabilities=frozenset({"resume", "cover_letter", "ats_safe", "docx"}),
    )

    formats: ClassVar[frozenset[str]] = DOCX_FORMATS

    async def render(
        self,
        doc: ResumeDocument,
        out: Path,
        *,
        fmt: str = DOCX_FORMAT,
        options: dict[str, Any] | None = None,
    ) -> RenderResult:
        """Build *doc* as a Word document at *out*, optionally converted to PDF.

        Args:
            doc: The resume to render.
            out: Destination path. Parent directories are created.
            fmt: ``"docx"`` or ``"pdf"``.
            options: ``font_size`` (points) and ``margin_in`` (inches) from
                ``render_resume``'s shrink loop, ``font`` (``"Calibri"`` or ``"Arial"``), and
                ``timeout`` for the LibreOffice conversion.

        Returns:
            The written file, its size, and its page count — measured from the PDF, or
            estimated from the document model for a ``.docx``, which has no fixed pagination
            until something lays it out.

        Raises:
            DocumentRenderError: If *fmt* is unsupported, python-docx is missing, the file
                cannot be written, or the PDF conversion fails.
        """
        require_format(self, fmt)
        out = self.resolve_output(out, fmt)
        values = self.merge_options(options)
        settings = self._settings_for(values)
        font = _resolve_font(values)
        document = build_resume_docx(doc, font=font, settings=settings)

        if fmt == DOCX_FORMAT:
            written = _save(document, out)
            pages = estimated_page_count(doc, settings)
            self.logger.info(
                "docx.rendered", path=str(out), bytes=written, font=font, estimated_pages=pages
            )
            return RenderResult.from_path(
                out, engine=DOCX_ENGINE, template=self.name, page_count=pages
            )

        await self._convert(document, out, values)
        result = RenderResult.from_path(out, engine=SOFFICE_ENGINE, template=self.name)
        self.logger.info(
            "docx.rendered_pdf",
            path=str(out),
            bytes=result.bytes_written,
            pages=result.page_count,
            font=font,
        )
        return result

    async def render_cover_letter(
        self,
        letter: CoverLetterDocument,
        out: Path,
        *,
        fmt: str = DOCX_FORMAT,
        options: dict[str, Any] | None = None,
    ) -> RenderResult:
        """Build *letter* as a Word document at *out*, optionally converted to PDF.

        Args:
            letter: The cover letter to render.
            out: Destination path. Parent directories are created.
            fmt: ``"docx"`` or ``"pdf"``.
            options: ``font_size``, ``margin_in``, ``font``, ``closing``, ``timeout``.

        Returns:
            The written file, its size, and its page count.

        Raises:
            DocumentRenderError: On the same conditions as :meth:`render`.
        """
        require_format(self, fmt)
        out = self.resolve_output(out, fmt)
        values = self.merge_options(options)
        settings = self._settings_for(values)
        font = _resolve_font(values)
        closing = str(values.get("closing") or DEFAULT_CLOSING)
        document = build_cover_letter_docx(letter, font=font, settings=settings, closing=closing)

        if fmt == DOCX_FORMAT:
            written = _save(document, out)
            self.logger.info("docx.cover_letter_rendered", path=str(out), bytes=written)
            return RenderResult.from_path(out, engine=DOCX_ENGINE, template=self.name, page_count=1)

        await self._convert(document, out, values)
        result = RenderResult.from_path(out, engine=SOFFICE_ENGINE, template=self.name)
        self.logger.info("docx.cover_letter_rendered_pdf", path=str(out), pages=result.page_count)
        return result

    async def _convert(
        self,
        document: Any,
        out: Path,
        options: dict[str, Any] | None,
    ) -> None:
        """Save a document to a scratch ``.docx`` and convert it to a PDF at *out*.

        The intermediate never lands next to the final PDF: a stray ``resume.docx`` beside
        ``resume.pdf`` in the user's documents folder is the kind of thing that gets attached
        to an application by mistake.

        Args:
            document: The built Word document.
            out: Destination PDF path.
            options: The caller's options, read for ``timeout``.

        Raises:
            DocumentRenderError: If the conversion fails for any reason.
        """
        timeout = float((options or {}).get("timeout") or SOFFICE_TIMEOUT_SECONDS)
        with tempfile.TemporaryDirectory(prefix="applicantos-docx-") as tmp:
            source = Path(tmp) / f"{out.stem or 'document'}.docx"
            _save(document, source)
            await docx_to_pdf(source, out, timeout=timeout)

    def _settings_for(self, options: dict[str, Any] | None) -> RenderSettings:
        """Resolve typography options, defaulting the margin to this template's half inch.

        Args:
            options: The caller's options mapping.

        Returns:
            The resolved settings.
        """
        merged = dict(options or {})
        merged.setdefault("margin_in", DEFAULT_DOCX_MARGIN_IN)
        return resolve_render_settings(merged)

    async def healthcheck(self) -> bool:
        """Report whether this template can build a ``.docx`` right now.

        Returns:
            ``True`` when python-docx imports. LibreOffice is deliberately *not* required:
            the template's primary format is ``docx``, which needs no external binary, and
            reporting the whole template unhealthy because no PDF converter is installed
            would take away the one renderer that always works on a bare Windows box.
        """
        try:
            _import_docx()
        except DocumentRenderError as exc:
            self.logger.warning("docx.unhealthy", error=str(exc))
            return False
        if not any(shutil.which(name) for name in SOFFICE_CANDIDATES):
            self.logger.debug("docx.soffice_absent", formats=["docx"])
        return True
