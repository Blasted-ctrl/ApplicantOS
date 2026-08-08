"""The generic document analyzer — and the file-to-text helpers every analyzer shares.

Two things live here, and they belong together.

**The shared reader.** Turning "a thing on disk or at a URL" into plain text is needed by
:class:`DocumentAnalyzer`, by :class:`~app.knowledge.analyzers.resume_parser.ResumeParser`
and by anything else that ever ingests a file. Implementing it twice would guarantee two
different answers for the same PDF, so it is implemented exactly once —
:func:`extract_text_from_path`, :func:`extract_text_from_bytes` and
:func:`fetch_text_from_url` — and imported by name from the modules that need it.

Every reader returns a :class:`TextExtraction` rather than raising: a PDF with no text
layer, a missing optional dependency and an undecodable byte soup are all *recoverable*
conditions that belong in :attr:`AnalysisResult.errors` where the user can see and fix
them. A resume that silently produces nothing is the single worst failure this subsystem
can have, so a scanned PDF says so in as many words.

**The last-resort analyzer.** :class:`DocumentAnalyzer` handles the source kinds that are
simply "some prose about the user" — a README, a documentation page, a blog post, an
interview note, a manual entry — and, because it also accepts any readable file path or
``http(s)`` URL, it is the fallback for a source kind no specialised analyzer claims. It
never *shadows* a specialised analyzer: it declares the most source kinds of any analyzer
in the package, and :func:`~app.knowledge.analyzers.base.analyzer_for` resolves the most
specific candidate first.

Third-party imports (``pypdf``, ``python-docx``, ``httpx``) are all lazy, inside the
function that needs them. ApplicantOS must import, start and index local Markdown on a
machine where none of them is installed.
"""

from __future__ import annotations

import codecs
import io
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import TYPE_CHECKING, Any, ClassVar, Final
from urllib.parse import unquote, urlparse
from urllib.request import url2pathname

import structlog

from app.knowledge.analyzers.base import (
    AnalysisResult,
    Analyzer,
    ExtractedDocument,
    SourceAccessDenied,
    SourceRef,
    SourceUnavailableError,
    compute_fingerprint,
    http_client,
)
from app.models.enums import FactKind, PluginKind, SourceKind
from app.plugins.base import PluginMeta
from app.plugins.registry import plugin

if TYPE_CHECKING:  # pragma: no cover - imported for annotations only
    from app.config.settings import Settings
    from app.knowledge.extractors import KnowledgeExtractor

__all__ = [
    "DEFAULT_FACT_KIND",
    "DOCX_SUFFIXES",
    "DocumentAnalyzer",
    "HTML_SUFFIXES",
    "MARKDOWN_SUFFIXES",
    "MAX_DOCUMENT_BYTES",
    "MIN_PDF_CHARS_PER_PAGE",
    "PDF_SUFFIXES",
    "PLAIN_TEXT_SUFFIXES",
    "SUPPORTED_DOCUMENT_SUFFIXES",
    "TextExtraction",
    "decode_bytes",
    "derive_title",
    "extract_text_from_bytes",
    "extract_text_from_path",
    "fetch_text_from_url",
    "file_probe_fingerprint",
    "knowledge_extractor",
    "local_path_for",
    "looks_like_url",
    "read_docx_text",
    "read_html_text",
    "read_pdf_text",
]

logger = structlog.get_logger(__name__)


# ======================================================================================
# Constants
# ======================================================================================

#: Suffixes routed to :func:`read_pdf_text`.
PDF_SUFFIXES: Final[frozenset[str]] = frozenset({".pdf"})

#: Suffixes routed to :func:`read_docx_text`. ``.doc`` is deliberately absent: the legacy
#: binary Word format needs a converter this project does not ship, and pretending to read
#: one would produce mojibake rather than an honest error.
DOCX_SUFFIXES: Final[frozenset[str]] = frozenset({".docx"})

#: Suffixes routed to :func:`read_html_text`.
HTML_SUFFIXES: Final[frozenset[str]] = frozenset({".html", ".htm", ".xhtml"})

#: Markdown-family suffixes, read as plain text (the extractors understand markdown).
MARKDOWN_SUFFIXES: Final[frozenset[str]] = frozenset({".md", ".markdown", ".mdx"})

#: Other prose suffixes read as plain text.
PLAIN_TEXT_SUFFIXES: Final[frozenset[str]] = frozenset(
    {".txt", ".text", ".rst", ".adoc", ".asciidoc", ".org", ".tex", ""}
)

#: Everything :func:`extract_text_from_path` knows how to read. Used by
#: :meth:`DocumentAnalyzer.supports` to decide whether it can act as a last resort for a
#: source kind nobody claimed.
SUPPORTED_DOCUMENT_SUFFIXES: Final[frozenset[str]] = (
    PDF_SUFFIXES | DOCX_SUFFIXES | HTML_SUFFIXES | MARKDOWN_SUFFIXES | PLAIN_TEXT_SUFFIXES
)

#: Largest file read into memory. A 32 MiB text document is already two orders of magnitude
#: larger than any resume or README; beyond this the file is data, not prose.
MAX_DOCUMENT_BYTES: Final[int] = 32 * 1024 * 1024

#: Below this many extracted characters per page, a PDF has no usable text layer — it is
#: almost always a scan or an image-only export. Reporting that is the whole point: a user
#: whose resume produced nothing needs to know *why*.
MIN_PDF_CHARS_PER_PAGE: Final[int] = 60

#: Bytes inspected when sniffing for binary content.
_BINARY_SNIFF_BYTES: Final[int] = 8_192

#: Text encodings tried, in order, by :func:`decode_bytes`. ``latin-1`` is last because it
#: cannot fail, which makes it the terminator rather than a real guess.
_TEXT_ENCODINGS: Final[tuple[str, ...]] = ("utf-8-sig", "cp1252", "latin-1")

#: Longest title derived from a document's own first line.
_MAX_TITLE_CHARS: Final[int] = 120

#: Default claim category for prose about the user that carries no stronger signal.
DEFAULT_FACT_KIND: Final[FactKind] = FactKind.ACCOMPLISHMENT

#: Domain-separation tags for the cheap change probes, so a file probe can never collide
#: with a URL probe or with a content fingerprint.
_FILE_PROBE_TAG: Final[str] = "analyzer.file.v1"
_URL_PROBE_TAG: Final[str] = "analyzer.url.v1"

#: HTTP statuses that mean "you need a credential" rather than "this is gone".
_FORBIDDEN_STATUSES: Final[frozenset[int]] = frozenset({401, 403, 407, 451})

#: Reader metadata not copied onto the emitted document. Outbound links belong to a crawler
#: deciding where to go next, not to a document that was fetched because the user named it.
_EXCLUDED_METADATA_KEYS: Final[frozenset[str]] = frozenset({"links"})

_MARKDOWN_H1: Final[re.Pattern[str]] = re.compile(r"^\s{0,3}#{1,6}\s+(?P<title>.+?)\s*#*\s*$")
_WHITESPACE_RUN: Final[re.Pattern[str]] = re.compile(r"\s+")


# ======================================================================================
# The reader's return type
# ======================================================================================


@dataclass(slots=True)
class TextExtraction:
    """Plain text recovered from a file, plus what was learned and what went wrong.

    Readers in this module never raise for content problems. A PDF with no text layer, a
    missing optional dependency, or bytes that decode to nothing all come back as an empty
    or partial :attr:`text` with an explanation in :attr:`errors`, which the analyzer
    forwards to :attr:`~app.knowledge.analyzers.base.AnalysisResult.errors` and the desktop
    app shows on the source.

    Attributes:
        text: The extracted plain text. May be empty.
        metadata: What the reader learned in passing — page count, table count, the
            extractor used, the HTTP content type.
        errors: Human-readable descriptions of recoverable problems.
    """

    text: str = ""
    metadata: dict[str, Any] = field(default_factory=dict)
    errors: list[str] = field(default_factory=list)

    @property
    def has_text(self) -> bool:
        """Whether any non-whitespace text was recovered."""
        return bool(self.text and self.text.strip())

    def record_error(self, message: str) -> None:
        """Append a recoverable problem description.

        Args:
            message: What went wrong, phrased for the person who owns the file.
        """
        text = str(message).strip()
        if text:
            self.errors.append(text)


# ======================================================================================
# Decoding
# ======================================================================================


def _is_probably_binary(data: bytes) -> bool:
    """Return whether *data* looks like a binary blob rather than text.

    Args:
        data: The raw bytes.

    Returns:
        ``True`` when a NUL byte appears in the sniffed prefix, which no text encoding this
        project reads ever produces mid-stream.
    """
    return b"\x00" in data[:_BINARY_SNIFF_BYTES]


def decode_bytes(data: bytes) -> str:
    """Decode text bytes, tolerating BOMs and the encodings resumes arrive in.

    UTF-16 is detected by its byte-order mark rather than guessed at, because decoding
    UTF-8 as UTF-16 succeeds and produces garbage. Everything else is tried as UTF-8 (BOM
    tolerated), then Windows-1252 — the encoding a Word "save as text" produces — then
    Latin-1, which cannot fail and therefore terminates the list.

    Args:
        data: Raw bytes believed to be text.

    Returns:
        The decoded string with newline conventions left as written; ``""`` for empty input.
    """
    if not data:
        return ""
    if data.startswith((codecs.BOM_UTF16_LE, codecs.BOM_UTF16_BE)):
        return data.decode("utf-16")
    if data.startswith((codecs.BOM_UTF32_LE, codecs.BOM_UTF32_BE)):
        return data.decode("utf-32")
    for encoding in _TEXT_ENCODINGS:
        try:
            return data.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return data.decode("utf-8", errors="replace")


# ======================================================================================
# PDF
# ======================================================================================


def read_pdf_text(data: bytes) -> TextExtraction:
    """Extract the text layer of a PDF.

    Uses ``pypdf``, imported lazily so the application runs without it. Encrypted files are
    opened with an empty password, which is what "protected against editing" PDFs — the
    common case for an exported resume — actually require.

    **A PDF that yields almost no text is reported, never swallowed.** Below
    :data:`MIN_PDF_CHARS_PER_PAGE` characters per page the file has no text layer at all,
    which means it is a scan or an image export; the caller gets that sentence in
    :attr:`TextExtraction.errors` instead of an inexplicably empty knowledge base.

    Args:
        data: The raw PDF bytes.

    Returns:
        The extracted text, with ``pages``/``extractor`` metadata and any problems found.
    """
    extraction = TextExtraction(metadata={"extractor": "pypdf"})
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        extraction.record_error(
            f"cannot read PDF: the 'pypdf' package is not installed ({exc}). "
            "Install it with `pip install pypdf`, or supply the document as .docx, .md or "
            ".txt."
        )
        return extraction

    try:
        reader = PdfReader(io.BytesIO(data))
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception as exc:  # noqa: BLE001 - any decrypt failure is the same story
                extraction.record_error(
                    f"PDF is password protected and could not be opened ({exc})."
                )
                return extraction
        pages = list(reader.pages)
    except Exception as exc:  # noqa: BLE001 - a malformed PDF must not kill an index run
        extraction.record_error(f"PDF could not be parsed ({type(exc).__name__}: {exc}).")
        return extraction

    parts: list[str] = []
    failed_pages = 0
    for number, page in enumerate(pages, start=1):
        try:
            parts.append(page.extract_text() or "")
        except Exception as exc:  # noqa: BLE001 - one bad page, not one bad document
            failed_pages += 1
            logger.debug("document.pdf_page_failed", page=number, error=str(exc))

    extraction.text = "\n\n".join(part for part in parts if part.strip())
    extraction.metadata["pages"] = len(pages)
    if failed_pages:
        extraction.metadata["failed_pages"] = failed_pages
        extraction.record_error(f"{failed_pages} of {len(pages)} PDF pages could not be read.")

    density = len(extraction.text.strip()) / len(pages) if pages else 0.0
    extraction.metadata["chars_per_page"] = round(density, 1)
    if pages and density < MIN_PDF_CHARS_PER_PAGE:
        extraction.metadata["likely_scanned"] = True
        extraction.record_error(
            f"PDF yielded only {len(extraction.text.strip())} characters across "
            f"{len(pages)} page(s): it is almost certainly a scanned image with no text "
            "layer. Re-export it from the original document, or run OCR over it, then "
            "re-index."
        )
    return extraction


# ======================================================================================
# DOCX
# ======================================================================================


def _docx_table_lines(table: Any) -> list[str]:
    """Render one ``python-docx`` table as text lines.

    Resumes are very often laid out in an invisible table — dates in the left column,
    content in the right — so a reader that only walks ``document.paragraphs`` returns an
    empty string for them. Rows whose cells are all single-line are joined with ``" | "``,
    which keeps a date beside the role it belongs to; a row containing a multi-line cell
    (a bullet list) is emitted cell by cell so those bullets survive as lines.

    Args:
        table: A ``docx.table.Table``.

    Returns:
        The rendered lines, in row order.
    """
    lines: list[str] = []
    for row in table.rows:
        try:
            cells = list(row.cells)
        except (IndexError, ValueError) as exc:  # malformed/merged grids
            logger.debug("document.docx_row_failed", error=str(exc))
            continue

        texts: list[str] = []
        for cell in cells:
            body = (cell.text or "").strip()
            for nested in getattr(cell, "tables", ()):
                body = "\n".join([body, *_docx_table_lines(nested)]).strip()
            # A vertically or horizontally merged cell is repeated across the span.
            if body and (not texts or texts[-1] != body):
                texts.append(body)
        if not texts:
            continue
        if any("\n" in text for text in texts):
            lines.extend(texts)
        else:
            lines.append(" | ".join(texts))
    return lines


def read_docx_text(data: bytes) -> TextExtraction:
    """Extract the text of a ``.docx`` file, including headers, tables and footers.

    Walks the document body in document order so paragraphs and tables interleave exactly
    as they are laid out, and reads each section's header and footer as well — a resume
    that puts the name and contact block in a Word header is common enough that ignoring
    it would lose the one part onboarding most wants.

    Args:
        data: The raw ``.docx`` bytes.

    Returns:
        The extracted text with ``paragraphs``/``tables`` metadata and any problems found.
    """
    extraction = TextExtraction(metadata={"extractor": "python-docx"})
    try:
        import docx
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError as exc:
        extraction.record_error(
            f"cannot read DOCX: the 'python-docx' package is not installed ({exc}). "
            "Install it with `pip install python-docx`, or supply the document as PDF, .md "
            "or .txt."
        )
        return extraction

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001 - a corrupt archive is a recoverable problem
        extraction.record_error(f"DOCX could not be parsed ({type(exc).__name__}: {exc}).")
        return extraction

    header_lines: list[str] = []
    footer_lines: list[str] = []
    try:
        for section in document.sections:
            for part, sink in ((section.header, header_lines), (section.footer, footer_lines)):
                sink.extend(
                    paragraph.text.strip()
                    for paragraph in part.paragraphs
                    if paragraph.text.strip()
                )
                for table in part.tables:
                    sink.extend(_docx_table_lines(table))
    except Exception as exc:  # noqa: BLE001 - headers are a bonus, never a requirement
        logger.debug("document.docx_headers_failed", error=str(exc))

    body_lines: list[str] = []
    paragraph_count = 0
    table_count = 0
    for child in document.element.body.iterchildren():
        tag = str(child.tag).rsplit("}", 1)[-1]
        if tag == "p":
            text = Paragraph(child, document).text.strip()
            paragraph_count += 1
            if text:
                body_lines.append(text)
        elif tag == "tbl":
            table_count += 1
            body_lines.extend(_docx_table_lines(Table(child, document)))

    ordered = [
        *dict.fromkeys(header_lines),
        *body_lines,
        *dict.fromkeys(footer_lines),
    ]
    extraction.text = "\n".join(ordered)
    extraction.metadata.update(
        {
            "paragraphs": paragraph_count,
            "tables": table_count,
            "header_lines": len(header_lines),
            "footer_lines": len(footer_lines),
        }
    )
    if not extraction.has_text:
        extraction.record_error(
            "DOCX contained no readable text (it may hold only images or drawing objects)."
        )
    return extraction


# ======================================================================================
# HTML
# ======================================================================================


def read_html_text(html: str, *, base_url: str | None = None) -> TextExtraction:
    """Convert an HTML document to readable plain text.

    Delegates to :func:`app.knowledge.analyzers._text.html_to_text`, which is the one HTML
    converter in the system — the website crawler and this module must agree on what a page
    "says", or the same portfolio would index differently depending on which analyzer saw
    it first. Its document metadata (``title``, ``description``, ``lang``, ``canonical``,
    ``headings``) is carried through onto :attr:`TextExtraction.metadata`; the title is
    additionally surfaced as ``html_title``, which is what an analyzer uses to name the
    document.

    Args:
        html: The raw HTML source.
        base_url: The URL the markup came from, so relative links resolve.

    Returns:
        The converted text and its metadata. A missing converter is recorded as an error
        rather than raised.
    """
    extraction = TextExtraction(metadata={"extractor": "html_to_text"})
    try:
        from app.knowledge.analyzers._text import html_to_text
    except ImportError as exc:  # pragma: no cover - only on a partially installed tree
        extraction.record_error(f"HTML conversion is unavailable ({exc}).")
        return extraction

    try:
        text, metadata = html_to_text(html or "", base_url=base_url)
    except Exception as exc:  # noqa: BLE001 - malformed markup is a content problem
        extraction.record_error(f"HTML could not be converted ({type(exc).__name__}: {exc}).")
        return extraction

    extraction.text = text
    extraction.metadata.update(metadata)
    title = metadata.get("title")
    if isinstance(title, str) and title.strip():
        extraction.metadata["html_title"] = _WHITESPACE_RUN.sub(" ", title).strip()[
            :_MAX_TITLE_CHARS
        ]
    return extraction


# ======================================================================================
# Dispatch
# ======================================================================================


def _looks_like_pdf(data: bytes) -> bool:
    """Return whether *data* begins with the PDF magic number."""
    return data[:5] == b"%PDF-"


def _looks_like_zip(data: bytes) -> bool:
    """Return whether *data* begins with a ZIP local-file header (a ``.docx`` container)."""
    return data[:4] == b"PK\x03\x04"


def extract_text_from_bytes(
    data: bytes,
    *,
    filename: str = "",
    content_type: str | None = None,
    base_url: str | None = None,
) -> TextExtraction:
    """Convert an in-memory document to plain text, choosing a reader by format.

    Format detection prefers evidence over declarations: the PDF and ZIP magic numbers beat
    both the filename suffix and a server's ``Content-Type``, because a portfolio that
    serves ``resume.pdf`` as ``application/octet-stream`` is entirely normal.

    Args:
        data: The raw document bytes.
        filename: Original filename or URL path, used for its suffix.
        content_type: HTTP ``Content-Type``, when the bytes came from a response.
        base_url: The URL the bytes came from, so relative links in HTML resolve.

    Returns:
        The extraction, including a ``format`` key in its metadata.
    """
    suffix = Path(filename or "").suffix.lower()
    declared = (content_type or "").split(";", 1)[0].strip().lower()

    if _looks_like_pdf(data) or suffix in PDF_SUFFIXES or "pdf" in declared:
        extraction = read_pdf_text(data)
        extraction.metadata.setdefault("format", "pdf")
    elif suffix in DOCX_SUFFIXES or (
        _looks_like_zip(data) and "wordprocessingml" in declared
    ):
        extraction = read_docx_text(data)
        extraction.metadata.setdefault("format", "docx")
    elif suffix in HTML_SUFFIXES or "html" in declared or "xml" in declared:
        extraction = read_html_text(decode_bytes(data), base_url=base_url)
        extraction.metadata.setdefault("format", "html")
    else:
        extraction = TextExtraction(metadata={"extractor": "plain", "format": "text"})
        if _is_probably_binary(data):
            extraction.record_error(
                f"{filename or 'document'} looks like a binary file, not text; no reader "
                "is registered for it."
            )
        else:
            extraction.text = decode_bytes(data)

    if content_type:
        extraction.metadata["content_type"] = content_type
    return extraction


def extract_text_from_path(path: Path) -> TextExtraction:
    """Read a local file and convert it to plain text.

    Args:
        path: An existing file. Directories and unreadable files produce an error rather
            than an exception, so one bad file never aborts a folder scan.

    Returns:
        The extraction, with ``size_bytes`` and ``modified_ns`` metadata.
    """
    extraction = TextExtraction()
    try:
        stat = path.stat()
    except OSError as exc:
        extraction.record_error(f"{path} could not be read ({exc.strerror or exc}).")
        return extraction

    if not path.is_file():
        extraction.record_error(f"{path} is not a file.")
        return extraction
    if stat.st_size > MAX_DOCUMENT_BYTES:
        extraction.record_error(
            f"{path.name} is {stat.st_size:,} bytes, above the "
            f"{MAX_DOCUMENT_BYTES:,}-byte document limit; skipped."
        )
        return extraction

    try:
        data = path.read_bytes()
    except OSError as exc:
        extraction.record_error(f"{path} could not be read ({exc.strerror or exc}).")
        return extraction

    extraction = extract_text_from_bytes(data, filename=path.name)
    extraction.metadata.update({"size_bytes": stat.st_size, "modified_ns": stat.st_mtime_ns})
    return extraction


async def fetch_text_from_url(url: str) -> TextExtraction:
    """Fetch a URL and convert the response body to plain text.

    Args:
        url: An absolute ``http`` or ``https`` URL.

    Returns:
        The extraction, with ``status_code``, ``etag`` and ``last_modified`` metadata when
        the server supplied them.

    Raises:
        SourceAccessDenied: If the server answered 401, 403, 407 or 451 — the remedy is a
            credential or a policy change, not a retry.
        SourceUnavailableError: If the URL could not be reached or returned any other error
            status. Nothing at all can be produced in that case.
    """
    extraction = TextExtraction()
    try:
        response = await http_client().get(url)
    except ImportError as exc:  # pragma: no cover - httpx absent on a slim install
        raise SourceUnavailableError(
            f"cannot fetch {url}: the 'httpx' package is not installed ({exc})."
        ) from exc
    except Exception as exc:  # noqa: BLE001 - httpx raises a family of transport errors
        raise SourceUnavailableError(
            f"{url} could not be fetched ({type(exc).__name__}: {exc})."
        ) from exc

    if response.status_code in _FORBIDDEN_STATUSES:
        raise SourceAccessDenied(
            f"{url} refused the request with HTTP {response.status_code}; it needs a "
            "credential this installation does not have."
        )
    if response.status_code >= 400:
        raise SourceUnavailableError(f"{url} returned HTTP {response.status_code}.")

    extraction = extract_text_from_bytes(
        response.content,
        filename=Path(unquote(urlparse(str(response.url)).path)).name,
        content_type=response.headers.get("content-type"),
        base_url=str(response.url),
    )
    extraction.metadata.update(
        {
            "status_code": response.status_code,
            "final_url": str(response.url),
            "etag": response.headers.get("etag"),
            "last_modified": response.headers.get("last-modified"),
        }
    )
    return extraction


# ======================================================================================
# Small shared helpers
# ======================================================================================


def looks_like_url(uri: str) -> bool:
    """Return whether *uri* is an absolute ``http``/``https`` URL.

    Args:
        uri: The source uri.

    Returns:
        ``True`` for a fetchable web address.
    """
    return uri.strip().lower().startswith(("http://", "https://"))


def local_path_for(uri: str) -> Path:
    """Return the filesystem path *uri* names, accepting ``file://`` URLs.

    Args:
        uri: An absolute or relative path, or a ``file://`` URL.

    Returns:
        The path, with ``~`` expanded. Not resolved and not checked for existence — callers
        that need a traversal guard must resolve it themselves.
    """
    candidate = uri.strip()
    if candidate.lower().startswith("file://"):
        parsed = urlparse(candidate)
        host = parsed.netloc
        path = url2pathname(parsed.path)
        candidate = rf"\\{host}{path}" if host and host != "localhost" else path
    return Path(candidate).expanduser()


def file_probe_fingerprint(path: Path) -> str:
    """Return a cheap change probe for one local file.

    Size and modification time answer "has this changed?" without reading a byte, which is
    the whole purpose of :meth:`~app.knowledge.analyzers.base.Analyzer.fingerprint`.
    Nanosecond precision matters: a file edited twice within the same second is a normal
    thing to do while iterating on a resume.

    Args:
        path: The file to probe.

    Returns:
        A 64-character digest, or ``""`` when the file cannot be stat'ed — which the caller
        must treat as "unknown", never as "unchanged".
    """
    try:
        stat = path.stat()
    except OSError:
        return ""
    return compute_fingerprint(_FILE_PROBE_TAG, path.as_posix(), stat.st_size, stat.st_mtime_ns)


def _url_probe_fingerprint(url: str, etag: str | None, last_modified: str | None) -> str:
    """Return a cheap change probe for one URL.

    Args:
        url: The address probed.
        etag: The server's ``ETag``, if any.
        last_modified: The server's ``Last-Modified``, if any.

    Returns:
        A digest over the validators, or ``""`` when the server offered none — a server
        that will not tell us whether a page changed must be re-read.
    """
    if not (etag or last_modified):
        return ""
    return compute_fingerprint(_URL_PROBE_TAG, url, etag or "", last_modified or "")


def derive_title(text: str, fallback: str) -> str:
    """Infer a human-facing title for a document from its own opening.

    Args:
        text: The document text.
        fallback: Title to use when the text does not announce one — normally the filename
            or the last URL segment.

    Returns:
        The first markdown heading, else the first non-empty line, else *fallback*.
    """
    for raw_line in (text or "").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        heading = _MARKDOWN_H1.match(line)
        candidate = heading.group("title") if heading else line
        cleaned = _WHITESPACE_RUN.sub(" ", candidate).strip(" \t#*_-–—:|")
        if cleaned:
            return cleaned[:_MAX_TITLE_CHARS]
    return fallback


def knowledge_extractor() -> KnowledgeExtractor:
    """Build the shared fact/entity extractor every analyzer in this package uses.

    Wires the application cache in so that re-indexing an unchanged document does not pay
    for a second model call (``docs/CONTRACTS.md`` §7 lists parsed resumes and project
    summaries as cache-by-contract). A cache that cannot be built is not an error: the
    extractor simply runs uncached, and — with no model configured at all — falls back to
    the fully deterministic rule-based path.

    Returns:
        A ready :class:`~app.knowledge.extractors.KnowledgeExtractor`.
    """
    from app.knowledge.extractors import KnowledgeExtractor

    cache = None
    try:
        from app.cache import get_cache

        cache = get_cache()
    except Exception as exc:  # noqa: BLE001 - an unavailable cache must not block indexing
        logger.debug("analyzer.cache_unavailable", error=str(exc))
    return KnowledgeExtractor(cache=cache)


# ======================================================================================
# The analyzer
# ======================================================================================


@plugin
class DocumentAnalyzer(Analyzer):
    """Indexes any single prose document — and is the fallback when nothing else fits.

    Handles the source kinds that are just "text about the user": a README, a documentation
    page, a blog post, an interview note, a manual entry. It additionally accepts any
    source whose uri is a readable local file or an ``http(s)`` URL, which is what makes it
    the last resort for a kind no specialised analyzer claims.

    It cannot shadow a specialised analyzer:
    :func:`~app.knowledge.analyzers.base.analyzer_for` sorts candidates by how few source
    kinds they declare, and this one declares the most in the package. A ``resume`` source
    therefore always resolves to
    :class:`~app.knowledge.analyzers.resume_parser.ResumeParser` even though this analyzer
    would happily read the same PDF.

    The pipeline is deliberately short — read, title, extract, fingerprint — because every
    interesting decision belongs to a component that is shared:
    :func:`extract_text_from_path` for the bytes,
    :class:`~app.knowledge.extractors.KnowledgeExtractor` for the knowledge.
    """

    meta: ClassVar[PluginMeta] = PluginMeta(
        kind=PluginKind.ANALYZER,
        name="document",
        version="1.0.0",
        display_name="Document",
        description=(
            "Reads a local file or a web page — Markdown, text, PDF, DOCX or HTML — and "
            "extracts facts, skills and projects from it."
        ),
        capabilities=frozenset({"local_file", "url", "pdf", "docx", "html", "fallback"}),
    )

    source_kinds: ClassVar[frozenset[SourceKind]] = frozenset(
        {
            SourceKind.README,
            SourceKind.DOCUMENTATION,
            SourceKind.BLOG_POST,
            SourceKind.INTERVIEW_NOTE,
            SourceKind.MANUAL_ENTRY,
        }
    )

    def __init__(self, settings: Settings, **kw: Any) -> None:
        """Construct the analyzer.

        Args:
            settings: Application settings, supplied by the plugin registry.
            **kw: Extra construction options, kept on :attr:`options`.
        """
        super().__init__(settings, **kw)
        self._extractor: KnowledgeExtractor | None = None

    # -- resolution -------------------------------------------------------------------

    def supports(self, source: SourceRef) -> bool:
        """Return whether this analyzer can read *source*.

        Args:
            source: The candidate source.

        Returns:
            ``True`` for a declared source kind, and — acting as the package's last resort
            — for any source whose uri is an ``http(s)`` URL or a file with a suffix
            :data:`SUPPORTED_DOCUMENT_SUFFIXES` covers. The filesystem is not touched:
            resolution must stay cheap and side-effect free.
        """
        if super().supports(source):
            return True
        uri = source.uri.strip()
        if not uri:
            return False
        if looks_like_url(uri):
            return True
        return Path(uri).suffix.lower() in SUPPORTED_DOCUMENT_SUFFIXES

    # -- change detection ---------------------------------------------------------------

    async def fingerprint(self, source: SourceRef) -> str:
        """Probe whether *source* changed, without reading its content.

        Args:
            source: The source to probe.

        Returns:
            A ``(path, size, mtime)`` digest for a local file, an ``ETag``/
            ``Last-Modified`` digest for a URL that offers validators, and the never-
            matching identity digest of
            :meth:`~app.knowledge.analyzers.base.Analyzer.fingerprint` when neither is
            available — so an unprobeable source is re-read rather than wrongly skipped.
        """
        uri = source.uri.strip()
        if uri and not looks_like_url(uri):
            probe = file_probe_fingerprint(local_path_for(uri))
            return probe or await super().fingerprint(source)

        if uri:
            try:
                response = await http_client().head(uri)
            except Exception as exc:  # noqa: BLE001 - a failed probe just means "unknown"
                logger.debug("document.head_probe_failed", uri=uri, error=str(exc))
            else:
                probe = _url_probe_fingerprint(
                    uri,
                    response.headers.get("etag"),
                    response.headers.get("last-modified"),
                )
                if probe:
                    return probe
        return await super().fingerprint(source)

    # -- analysis -----------------------------------------------------------------------

    async def analyze(self, source: SourceRef) -> AnalysisResult:
        """Read *source*, extract its knowledge, and return it.

        Args:
            source: A local file, a ``file://`` URL, or an ``http(s)`` URL.

        Returns:
            One document plus the facts, entities and edges the extractor could ground in
            it. A document that could not be read comes back empty with the reason in
            :attr:`~app.knowledge.analyzers.base.AnalysisResult.errors`.

        Raises:
            SourceUnavailableError: If the file does not exist or the URL is unreachable.
            SourceAccessDenied: If the file or URL cannot be read for permission reasons.
        """
        self.require_supported(source)
        result = AnalysisResult()
        uri = source.uri.strip()
        if not uri:
            raise SourceUnavailableError("document source has no uri", source=source)

        if looks_like_url(uri):
            extraction = await fetch_text_from_url(uri)
            canonical_uri = str(extraction.metadata.get("final_url") or uri)
            fallback_title = Path(unquote(urlparse(canonical_uri).path)).name or canonical_uri
            probe = _url_probe_fingerprint(
                uri,
                extraction.metadata.get("etag"),
                extraction.metadata.get("last_modified"),
            )
        else:
            path = local_path_for(uri)
            try:
                resolved = path.resolve(strict=True)
            except OSError as exc:
                raise SourceUnavailableError(
                    f"{path} does not exist or cannot be resolved ({exc.strerror or exc})",
                    source=source,
                ) from exc
            if not resolved.is_file():
                raise SourceUnavailableError(f"{resolved} is not a file", source=source)
            extraction = extract_text_from_path(resolved)
            canonical_uri = resolved.as_posix()
            fallback_title = resolved.name
            probe = file_probe_fingerprint(resolved)

        for message in extraction.errors:
            result.record_error(message)

        if not extraction.has_text:
            result.fingerprint = probe or compute_fingerprint(_FILE_PROBE_TAG, canonical_uri, 0)
            logger.info("document.no_text", uri=canonical_uri, errors=len(result.errors))
            return result

        title = (
            source.label
            or str(extraction.metadata.get("html_title") or "")
            or derive_title(extraction.text, fallback_title)
        )
        document = ExtractedDocument(
            uri=canonical_uri,
            title=title,
            text=extraction.text,
            kind=source.kind,
            metadata={
                "analyzer": self.name,
                "source_uri": uri,
                **{
                    key: value
                    for key, value in extraction.metadata.items()
                    if value is not None and key not in _EXCLUDED_METADATA_KEYS
                },
            },
        )
        result.documents.append(document)

        extracted = await self._knowledge().extract(
            extraction.text,
            kind=DEFAULT_FACT_KIND,
            context={
                "organization": source.option("organization"),
                "role": source.option("role"),
                "source_uri": canonical_uri,
            },
        )
        result.merge(extracted).deduplicate()
        result.fingerprint = probe or result.content_fingerprint()

        logger.info("document.analyzed", uri=canonical_uri, **result.counts())
        return result

    # -- internals ----------------------------------------------------------------------

    def _knowledge(self) -> KnowledgeExtractor:
        """Return this instance's extractor, building it once.

        Returns:
            The shared-cache-backed :class:`~app.knowledge.extractors.KnowledgeExtractor`.
        """
        if self._extractor is None:
            self._extractor = knowledge_extractor()
        return self._extractor
